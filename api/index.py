from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
import os
import tempfile
import img2pdf
import zipfile
from io import BytesIO
import uuid

app = Flask(__name__, template_folder='../templates', static_folder='../static')

# Use /tmp for Vercel serverless functions
app.config['UPLOAD_FOLDER'] = '/tmp'
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

def cleanup(path):
    try:
        if os.path.exists(path):
            os.remove(path)
    except Exception as e:
        print(f"Cleanup failed: {e}")


def validate_page_range(pages_input, total_pages):
    # Validate if the page range is valid for the given PDF
    if not pages_input.strip():
        return True, None
    
    try:
        invalid_pages = []
        for part in pages_input.split(','):
            part = part.strip()
            if '-' in part:
                start, end = map(int, part.split('-'))
                if start < 1 or end < 1 or start > total_pages or end > total_pages or start > end:
                    invalid_pages.extend([str(p) for p in range(start, end + 1) if p > total_pages or p < 1])
            else:
                page = int(part)
                if page < 1 or page > total_pages:
                    invalid_pages.append(str(page))
        
        if invalid_pages:
            return False, f"Error: Pages {', '.join(set(invalid_pages))} do not exist. PDF has only {total_pages} pages."
        return True, None
    except ValueError:
        return False, "Error: Invalid page format. Use format like: 1-3,5,7"

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/merge', methods=['POST'])
def merge_pdfs():
    file_inputs = []
    i = 0
    while True:
        f = request.files.get(f'file_{i}')
        if f is None or f.filename == '':
            break
        file_inputs.append((f, i))
        i += 1
    if not file_inputs:
        return "No files uploaded", 400
    merger = PdfMerger()
    temp_files = []
    try:
        for file, idx in file_inputs:
            filename = secure_filename(file.filename)
            if not filename.lower().endswith('.pdf'):
                continue
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            temp_files.append(file_path)
            pages_input = request.form.get(f'pages_{idx}', '').strip()
            if pages_input:
                # Validate pages for this PDF
                temp_reader = PdfReader(file_path)
                total_pages = len(temp_reader.pages)
                is_valid, error_msg = validate_page_range(pages_input, total_pages)
                if not is_valid:
                    for fp in temp_files:
                        cleanup(fp)
                    return f"{error_msg} (File: {filename})", 400
                
                for part in pages_input.split(','):
                    part = part.strip()
                    if '-' in part:
                        try:
                            start, end = map(int, part.split('-'))
                            merger.append(file_path, pages=(start-1, end))
                        except:
                            pass
                    else:
                        try:
                            page = int(part)
                            merger.append(file_path, pages=(page-1, page))
                        except:
                            pass
            else:
                merger.append(file_path)
        if len(merger.pages) == 0:
            return "No pages to merge", 400
        output_filename = f"merged_{uuid.uuid4().hex}.pdf"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        merger.write(output_path)
        merger.close()
        response = send_file(output_path, as_attachment=True, download_name="merged.pdf", mimetype='application/pdf')
        @response.call_on_close
        def remove_files():
            cleanup(output_path)
            for fp in temp_files:
                cleanup(fp)
        return response
    except Exception as e:
        return f"Error: {str(e)}", 500

@app.route('/split', methods=['POST'])
def split_pdf():
    if 'file' not in request.files or not request.files['file'].filename:
        return "Error: No file uploaded", 400
    pages_input = request.form.get('pages', '').strip()
    if not pages_input:
        return "Error: No pages specified", 400
    try:
        file = request.files['file']
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        # Validate page range
        is_valid, error_msg = validate_page_range(pages_input, total_pages)
        if not is_valid:
            cleanup(input_path)
            return error_msg, 400
        
        writer = PdfWriter()
        for part in pages_input.split(','):
            part = part.strip()
            if '-' in part:
                start, end = map(int, part.split('-'))
                for p in range(start-1, end):
                    if 0 <= p < len(reader.pages):
                        writer.add_page(reader.pages[p])
            else:
                p = int(part) - 1
                if 0 <= p < len(reader.pages):
                    writer.add_page(reader.pages[p])
        if not writer.pages:
            cleanup(input_path)
            return "Error: No valid pages selected", 400
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"split_{uuid.uuid4().hex}.pdf")
        with open(output_path, 'wb') as f:
            writer.write(f)
        response = send_file(output_path, as_attachment=True, download_name="split.pdf", mimetype='application/pdf')
        @response.call_on_close
        def cleanup_files():
            cleanup(input_path)
            cleanup(output_path)
        return response
    except Exception as e:
        return f"Error: {str(e)}", 500

@app.route('/compress', methods=['POST'])
def compress_pdf():
    if 'file' not in request.files or not request.files['file'].filename:
        return "No file uploaded", 400
    try:
        import fitz
        file = request.files['file']
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        doc = fitz.open(input_path)
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"compressed_{uuid.uuid4().hex}.pdf")
        doc.save(output_path, garbage=4, deflate=True, clean=True, ascii=False, pretty=False, linear=False, encryption=0)
        doc.close()
        response = send_file(output_path, as_attachment=True, download_name="compressed.pdf", mimetype='application/pdf')
        @response.call_on_close
        def cleanup_files():
            cleanup(input_path)
            cleanup(output_path)
        return response
    except Exception as e:
        return f"Compression error: {str(e)}", 500

@app.route('/pdf_to_word', methods=['POST'])
def pdf_to_word():
    if 'file' not in request.files or not request.files['file'].filename:
        return "No file uploaded", 400
    try:
        import fitz
        from docx import Document
        from docx.shared import Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        file = request.files['file']
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"word_{uuid.uuid4().hex}.docx")
        doc = fitz.open(input_path)
        document = Document()
        section = document.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes("png")
            img_stream = io.BytesIO(img_data)
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(img_stream, width=Inches(6.5))
            if page_num < len(doc) - 1:
                document.add_page_break()
        doc.close()
        document.save(output_path)
        response = send_file(output_path, as_attachment=True, download_name="document.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        @response.call_on_close
        def cleanup_files():
            cleanup(input_path)
            cleanup(output_path)
        return response
    except Exception as e:
        return f"Conversion error: {str(e)}", 500

@app.route('/pdf_to_images', methods=['POST'])
def pdf_to_images():
    if 'file' not in request.files or not request.files['file'].filename:
        return "Error: No file uploaded", 400
    pages_input = request.form.get('pages', '').strip()
    if not pages_input:
        return "Error: No pages specified", 400
    try:
        import fitz
        file = request.files['file']
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)
        doc = fitz.open(input_path)
        total_pages = len(doc)
        
        # Validate page range
        is_valid, error_msg = validate_page_range(pages_input, total_pages)
        if not is_valid:
            doc.close()
            cleanup(input_path)
            return error_msg, 400
        
        page_numbers = []
        for part in pages_input.split(','):
            part = part.strip()
            if '-' in part:
                start, end = map(int, part.split('-'))
                page_numbers.extend(range(start-1, end))
            else:
                page_numbers.append(int(part)-1)
        page_numbers = [p for p in page_numbers if 0 <= p < len(doc)]
        if not page_numbers:
            doc.close()
            cleanup(input_path)
            return "Error: No valid pages selected", 400
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zipf:
            for p in page_numbers:
                page = doc.load_page(p)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                zipf.writestr(f"page_{p+1}.png", img_data)
        doc.close()
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"pdf_to_images_{uuid.uuid4().hex}.zip")
        with open(output_path, 'wb') as f:
            f.write(zip_buffer.getvalue())
        response = send_file(output_path, as_attachment=True, download_name="pages.zip", mimetype='application/zip')
        @response.call_on_close
        def cleanup_files():
            cleanup(input_path)
            cleanup(output_path)
        return response
    except Exception as e:
        return f"Error: {str(e)}", 500

@app.route('/images_to_pdf', methods=['POST'])
def images_to_pdf():
    files = request.files.getlist('file')
    if not files or all(f.filename == '' for f in files):
        return "No images uploaded", 400
    try:
        images = []
        temp_paths = []
        for file in files:
            if file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.JPG', '.JPEG')):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                temp_paths.append(file_path)
                with open(file_path, 'rb') as img_file:
                    images.append(img_file.read())
        if not images:
            return "No valid images found", 400
        pdf_bytes = img2pdf.convert(images)
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"images_to_pdf_{uuid.uuid4().hex}.pdf")
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)
        response = send_file(output_path, as_attachment=True, download_name="images.pdf", mimetype='application/pdf')
        @response.call_on_close
        def cleanup_files():
            cleanup(output_path)
            for path in temp_paths:
                cleanup(path)
        return response
    except Exception as e:
        return f"Error: {str(e)}", 500

# Export app for Vercel

if __name__ == '__main__':
    print("Starting PDF Toolbox...")
    app.run(debug=True)
