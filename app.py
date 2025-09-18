from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
import os
import tempfile
import img2pdf
from pdf2docx import Converter
from pdf2image import convert_from_path
import zipfile
from io import BytesIO
import uuid

app = Flask(__name__)

app.config['UPLOAD_FOLDER'] = tempfile.mkdtemp()

app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB

def cleanup(path):
    try:
        if os.path.exists(path):  # Cek apakah file masih ada
            os.remove(path)  # Hapus file
    except Exception as e:
        print(f"Cleanup failed: {e}")  # Tampilkan error jika gagal


@app.route('/')
def index():
    return render_template('index.html')


# --- MERGE PDFs ---
@app.route('/merge', methods=['POST'])
def merge_pdfs():
    # Simpan semua file yang diunggah ke dalam list
    file_inputs = []
    i = 0  # Mulai dari file ke-0
    while True:
        # Ambil file dengan nama 'file_0', 'file_1', dst.
        f = request.files.get(f'file_{i}')
        # Jika tidak ada file atau nama kosong, hentikan loop
        if f is None or f.filename == '':
            break
        # Tambahkan file dan indeksnya ke list
        file_inputs.append((f, i))
        i += 1  # Naikkan indeks

    # Jika tidak ada file yang diunggah, kirim error
    if not file_inputs:
        return "No files uploaded", 400

    # Siapkan objek untuk menggabungkan PDF
    merger = PdfMerger()
    temp_files = []  # Untuk menyimpan path file sementara

    try:
        # Proses setiap file yang diunggah
        for file, idx in file_inputs:
            filename = secure_filename(file.filename)  # Nama file yang aman
            # Lewati jika bukan file PDF
            if not filename.lower().endswith('.pdf'):
                continue

            # Simpan file ke folder sementara
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            temp_files.append(file_path)  # Simpan path untuk nanti dihapus

            # Ambil input halaman dari form (misal: "1-3,5")
            pages_input = request.form.get(f'pages_{idx}', '').strip()
            if pages_input:
                # Pisahkan per bagian (misal: "1-3" dan "5")
                for part in pages_input.split(','):
                    part = part.strip()
                    if '-' in part:
                        # Jika bentuknya "1-3" → ambil dari halaman 1 sampai 3
                        try:
                            start, end = map(int, part.split('-'))
                            merger.append(file_path, pages=(start-1, end))  # PyPDF2: 0-indexed
                        except:
                            pass  # Lewati jika format salah
                    else:
                        # Jika satu halaman, misal "5"
                        try:
                            page = int(part)
                            merger.append(file_path, pages=(page-1, page))
                        except:
                            pass  # Lewati jika format salah
            else:
                # Jika tidak ada input halaman → gabung semua halaman
                merger.append(file_path)

        # Jika tidak ada halaman yang berhasil digabung
        if len(merger.pages) == 0:
            return "No pages to merge", 400

        # Buat nama file output yang unik
        output_filename = f"merged_{uuid.uuid4().hex}.pdf"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        merger.write(output_path)  # Simpan hasil gabungan
        merger.close()

        # Kirim file hasil ke pengguna sebagai download
        response = send_file(
            output_path,
            as_attachment=True,
            download_name="merged.pdf"
        )

        # Hapus file sementara setelah download selesai
        @response.call_on_close
        def remove_files():
            cleanup(output_path)  # Hapus file hasil
            for fp in temp_files:
                cleanup(fp)  # Hapus file input

        return response

    except Exception as e:
        return f"Error: {str(e)}", 500  # Tampilkan error jika terjadi


# --- SPLIT PDF ---
@app.route('/split', methods=['POST'])
def split_pdf():
    """
    Memotong halaman tertentu dari satu PDF dan menyimpannya sebagai file baru.
    Input: file PDF dan daftar halaman (misal: "1-3,5")
    """
    # Cek apakah file diunggah
    if 'file' not in request.files or not request.files['file'].filename:
        return "No file uploaded", 400

    # Ambil input halaman dari form
    pages_input = request.form.get('pages', '').strip()
    if not pages_input:
        return "No pages specified", 400

    try:
        file = request.files['file']
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)  # Simpan file sementara

        reader = PdfReader(input_path)
        writer = PdfWriter()  # Penulis PDF baru

        # Proses setiap bagian input halaman
        for part in pages_input.split(','):
            part = part.strip()
            if '-' in part:
                # Range halaman: "1-3" → halaman 1, 2, 3
                start, end = map(int, part.split('-'))
                for p in range(start-1, end):  # 0-indexed
                    if 0 <= p < len(reader.pages):
                        writer.add_page(reader.pages[p])
            else:
                # Satu halaman: "5"
                p = int(part) - 1
                if 0 <= p < len(reader.pages):
                    writer.add_page(reader.pages[p])

        # Jika tidak ada halaman yang dipilih
        if not writer.pages:
            return "No valid pages selected", 400

        # Simpan hasil split
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"split_{uuid.uuid4().hex}.pdf")
        with open(output_path, 'wb') as f:
            writer.write(f)

        # Kirim file ke pengguna
        response = send_file(
            output_path,
            as_attachment=True,
            download_name="split.pdf"
        )

        # Hapus file sementara setelah download
        @response.call_on_close
        def cleanup_files():
            cleanup(input_path)
            cleanup(output_path)

        return response

    except Exception as e:
        return f"Error: {str(e)}", 500


# --- COMPRESS PDF ---
@app.route('/compress', methods=['POST'])
def compress_pdf():
    """
    Menyusutkan ukuran file PDF dengan mengoptimalkan struktur internal.
    Tidak mengurangi kualitas gambar, hanya mengompres struktur file.
    """
    if 'file' not in request.files or not request.files['file'].filename:
        return "No file uploaded", 400

    try:
        import fitz  # PyMuPDF — library canggih untuk PDF

        file = request.files['file']
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        # Buka PDF
        doc = fitz.open(input_path)
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"compressed_{uuid.uuid4().hex}.pdf")

        # Simpan dengan kompresi tinggi
        doc.save(
            output_path,
            garbage=4,        # Bersihkan objek yang tidak terpakai
            deflate=True,     # Kompresi data
            clean=True,       # Optimalkan struktur
            ascii=False,      # Jangan encode ke ASCII
            pretty=False,     # Jangan format rapi (untuk ukuran kecil)
            linear=False,     # Jangan linearize (save space)
            encryption=0      # Tidak ada enkripsi
        )
        doc.close()

        # Kirim file hasil
        response = send_file(
            output_path,
            as_attachment=True,
            download_name="compressed.pdf"
        )

        # Hapus file setelah download
        @response.call_on_close
        def cleanup_files():
            cleanup(input_path)
            cleanup(output_path)

        return response

    except Exception as e:
        return f"Compression error: {str(e)}", 500


# --- PDF TO WORD ---
@app.route('/pdf_to_word', methods=['POST'])
def pdf_to_word():
    """
    Mengubah PDF menjadi dokumen Word (.docx).
    Setiap halaman PDF diubah menjadi gambar, lalu dimasukkan ke Word.
    Cocok untuk PDF yang tidak bisa diekstrak teksnya.
    """
    if 'file' not in request.files or not request.files['file'].filename:
        return "No file uploaded", 400

    try:
        import fitz  # Untuk baca PDF
        from docx import Document  # Untuk buat dokumen Word
        from docx.shared import Inches, Cm  # Ukuran
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # Rata tengah
        import io  # Untuk handle data biner

        file = request.files['file']
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        # Buat dokumen Word baru
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"word_{uuid.uuid4().hex}.docx")
        doc = fitz.open(input_path)
        document = Document()

        # Atur ukuran halaman A4
        section = document.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

        # Konversi setiap halaman PDF ke gambar
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            mat = fitz.Matrix(2.0, 2.0)  # Zoom 2x untuk kualitas tinggi
            pix = page.get_pixmap(matrix=mat, alpha=False)  # PNG tanpa transparansi
            img_data = pix.tobytes("png")
            img_stream = io.BytesIO(img_data)

            # Tambahkan gambar ke dokumen Word
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(img_stream, width=Inches(6.5))  # Lebar 6.5 inch

            # Tambahkan page break kecuali di halaman terakhir
            if page_num < len(doc) - 1:
                document.add_page_break()

        doc.close()
        document.save(output_path)  # Simpan dokumen

        # Kirim ke pengguna
        response = send_file(
            output_path,
            as_attachment=True,
            download_name="document.docx"
        )

        # Hapus file setelah download
        @response.call_on_close
        def cleanup_files():
            cleanup(input_path)
            cleanup(output_path)

        return response

    except Exception as e:
        return f"Conversion error: {str(e)}", 500


# --- PDF TO IMAGES ---
@app.route('/pdf_to_images', methods=['POST'])
def pdf_to_images():
    """
    Mengubah halaman tertentu dari PDF menjadi gambar (PNG).
    Hasilnya dikompres dalam file ZIP.
    """
    if 'file' not in request.files or not request.files['file'].filename:
        return "No file uploaded", 400

    pages_input = request.form.get('pages', '').strip()
    if not pages_input:
        return "No pages specified", 400

    try:
        import fitz
        file = request.files['file']
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        doc = fitz.open(input_path)
        page_numbers = []

        # Parse input halaman (misal: "1-3,5")
        for part in pages_input.split(','):
            part = part.strip()
            if '-' in part:
                start, end = map(int, part.split('-'))
                page_numbers.extend(range(start-1, end))  # 0-indexed
            else:
                page_numbers.append(int(part)-1)
        # Filter halaman yang valid
        page_numbers = [p for p in page_numbers if 0 <= p < len(doc)]

        if not page_numbers:
            return "No valid pages", 400

        # Buat ZIP berisi gambar
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zipf:
            for p in page_numbers:
                page = doc.load_page(p)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                zipf.writestr(f"page_{p+1}.png", img_data)
        doc.close()

        # Simpan ZIP sementara
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"pdf_to_images_{uuid.uuid4().hex}.zip")
        with open(output_path, 'wb') as f:
            f.write(zip_buffer.getvalue())

        # Kirim ZIP
        response = send_file(
            output_path,
            as_attachment=True,
            download_name="pages.zip"
        )

        # Hapus file setelah download
        @response.call_on_close
        def cleanup_files():
            cleanup(input_path)
            cleanup(output_path)

        return response

    except Exception as e:
        return f"Error: {str(e)}", 500


# --- IMAGES TO PDF ---
@app.route('/images_to_pdf', methods=['POST'])
def images_to_pdf():
    """
    Menggabungkan beberapa gambar menjadi satu file PDF.
    Mendukung format: PNG, JPG, JPEG.
    """
    files = request.files.getlist('file')
    if not files or all(f.filename == '' for f in files):
        return "No images uploaded", 400

    try:
        images = []
        temp_paths = []

        # Proses setiap gambar
        for file in files:
            if file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.JPG', '.JPEG')):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                temp_paths.append(file_path)
                with open(file_path, 'rb') as img_file:
                    images.append(img_file.read())

        if not images:
            return "No valid images found", 400

        # Konversi gambar ke PDF
        pdf_bytes = img2pdf.convert(images)
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"images_to_pdf_{uuid.uuid4().hex}.pdf")
        with open(output_path, 'wb') as f:
            f.write(pdf_bytes)

        # Kirim hasil
        response = send_file(
            output_path,
            as_attachment=True,
            download_name="images.pdf"
        )

        # Hapus file sementara
        @response.call_on_close
        def cleanup_files():
            cleanup(output_path)
            for path in temp_paths:
                cleanup(path)

        return response

    except Exception as e:
        return f"Error: {str(e)}", 500  


# Jalankan aplikasi
if __name__ == '__main__':
    print("Starting PDF Toolbox...")
    app.run(debug=True)